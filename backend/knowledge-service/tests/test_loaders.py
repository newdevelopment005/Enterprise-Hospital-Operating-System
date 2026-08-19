"""Tests for the Medical Knowledge Base document loaders."""

import io
import json

import pytest

from knowledge_service.service.loaders import LoaderError, load_documents


def _make_pdf(text: str) -> bytes:
    from reportlab.pdfgen import canvas

    stream = io.BytesIO()
    c = canvas.Canvas(stream)
    c.drawString(72, 720, text[:90])
    c.showPage()
    c.save()
    return stream.getvalue()


def _make_docx(text: str, table: tuple = ()) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("Medication Administration Policy", 0)
    document.add_paragraph(text)
    if table:
        t = document.add_table(rows=len(table), cols=2)
        for i, (left, right) in enumerate(table):
            t.cell(i, 0).text = left
            t.cell(i, 1).text = right
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_pdf_loader_extracts_text(settings):
    raw = _make_pdf("Aspirin 300 mg should be given immediately in STEMI.")
    docs = load_documents(raw, filename="ami-guideline.pdf", settings=settings)
    assert len(docs) == 1
    assert docs[0].source_format == "PDF"
    assert docs[0].doc_type == "GUIDELINE"
    assert "Aspirin 300 mg" in docs[0].text
    assert docs[0].ingestion_ref == "ami-guideline.pdf"


def test_word_loader_keeps_headings_and_tables(settings):
    raw = _make_docx(
        "Verify two identifiers before every dose.",
        (("Right patient", "name and DOB"), ("Right drug", "check against order")),
    )
    docs = load_documents(raw, filename="policy.docx", settings=settings)
    assert len(docs) == 1
    assert docs[0].source_format == "DOCX"
    assert "# Medication Administration Policy" in docs[0].text
    assert "|" in docs[0].text
    assert "Right drug" in docs[0].text


def test_sop_loader_normalizes_sections_and_steps(settings):
    sop = {
        "title": "Central Line Insertion SOP",
        "department": "ICU",
        "owner": "Nursing Director",
        "purpose": "Standardize sterile insertion.",
        "sections": [
            {"title": "Procedure", "steps": ["Hand hygiene", "Full barrier", "Ultrasound guidance"]}
        ],
    }
    docs = load_documents(json.dumps(sop).encode(), filename="sop.json", kind="sop", settings=settings)
    assert len(docs) == 1
    assert docs[0].doc_type == "PROTOCOL"
    assert "ICU" in docs[0].text
    assert "1. Hand hygiene" in docs[0].text
    assert "Full barrier" in docs[0].text
    assert docs[0].metadata["department"] == "ICU"


def test_formulary_csv_produces_per_drug_documents(settings):
    csv_text = (
        "drug_name,generic_name,class,indications,contraindications,interactions\n"
        "Paracetamol,acetaminophen,analgesic,fever,pain,liver failure,\n"
        "Aspirin,acetylsalicylic acid,NSAID,STEMI,allergy,bleeding risk\n"
    )
    docs = load_documents(
        csv_text.encode(), filename="formulary.csv", doc_type="MEDICATION", settings=settings
    )
    assert len(docs) == 2
    names = {d.title for d in docs}
    assert {"Paracetamol", "Aspirin"} == names
    paracetamol = next(d for d in docs if d.title == "Paracetamol")
    assert "fever" in paracetamol.text
    assert paracetamol.metadata["generic_name"] == "acetaminophen"


def test_formulary_json_requires_formulary_doc_type(settings):
    payload = json.dumps(
        [{"drug_name": "Amoxicillin", "generic_name": "amoxicillin", "class": "beta-lactam"}]
    ).encode()
    docs = load_documents(payload, filename="formulary.json", doc_type="FORMULARY", settings=settings)
    assert len(docs) == 1
    assert "Amoxicillin" in docs[0].text


def test_book_loader_chapters_are_separate_documents(settings):
    book = {
        "title": "Internal Medicine Essentials",
        "author": "A. Physician",
        "edition": "2",
        "chapters": [
            {"number": 1, "title": "Cardiology", "content": "Heart failure management basics."},
            {"number": 2, "title": "Nephrology", "content": "Acute kidney injury approach."},
        ],
    }
    docs = load_documents(
        json.dumps(book).encode(), filename="medicine.book.json", settings=settings
    )
    assert len(docs) == 2
    assert docs[0].doc_type == "TEXTBOOK"
    assert all(d.metadata["book"] == "Internal Medicine Essentials" for d in docs)
    assert any("Cardiology" in d.title for d in docs)
    assert any("Nephrology" in d.title for d in docs)


def test_journal_loader_articles_are_separate_documents(settings):
    issue = {
        "journal_name": "Hospital Medicine Journal",
        "volume": "12",
        "issue": "3",
        "articles": [
            {"title": "Sepsis bundles", "authors": ["B. Clinician"], "abstract": "Early lactate.",
             "keywords": ["sepsis"], "body": "Bundle compliance improves survival."},
        ],
    }
    docs = load_documents(
        json.dumps(issue).encode(), filename="issue.journal.json", settings=settings
    )
    assert len(docs) == 1
    article = docs[0]
    assert article.doc_type == "JOURNAL"
    assert article.source_format == "JOURNAL"
    assert "Journal: Hospital Medicine Journal | Volume: 12 | Issue: 3" in article.text
    assert "Bundle compliance" in article.text
    assert article.metadata["doi"] is None
    assert article.metadata["keywords"] == ["sepsis"]


def test_markdown_and_text_loaders_fall_back(settings):
    md = load_documents(b"# Hand Hygiene Policy\nWash hands for 20 seconds.", filename="hygiene.md", settings=settings)
    assert md[0].source_format == "MARKDOWN"
    assert md[0].title == "Hand Hygiene Policy"
    txt = load_documents(b"Store cold chain vaccines at 2-8 C.", filename="vaccine.txt", settings=settings)
    assert txt[0].source_format == "TEXT"


def test_unknown_loader_kind_rejected(settings):
    with pytest.raises(LoaderError) as exc:
        load_documents(b"{}", filename="x.json", kind="nope", settings=settings)
    assert exc.value.code == "UNSUPPORTED_FORMAT"


def test_ambiguous_json_requires_kind(settings):
    with pytest.raises(LoaderError) as exc:
        load_documents(b"{}", filename="x.json", settings=settings)
    assert exc.value.code == "UNSUPPORTED_FORMAT"
    assert "kind" in exc.value.message


def test_scanned_pdf_without_text_rejected(settings):
    from reportlab.pdfgen import canvas

    stream = io.BytesIO()
    c = canvas.Canvas(stream)
    c.showPage()
    c.save()
    with pytest.raises(LoaderError) as exc:
        load_documents(stream.getvalue(), filename="scan.pdf", settings=settings)
    assert exc.value.code == "PARSE_ERROR"