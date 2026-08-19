"""Word (.docx) loader — headings + paragraphs + tables via python-docx."""

from __future__ import annotations

import io
import os

from knowledge_service.service.loaders.base import DocumentLoader, LoadedDocument, LoaderError


class WordLoader(DocumentLoader):
    """Extract a .docx document as lightweight Markdown (headings -> # levels)."""

    formats = (".docx",)
    kind = "word"

    def load(self, raw: bytes, *, filename: str) -> list[LoadedDocument]:
        try:
            from docx import Document  # type: ignore[import-not-found]
        except ImportError as err:  # pragma: no cover - optional dependency
            raise LoaderError(
                "LOADER_UNAVAILABLE",
                "Word parsing requires the 'python-docx' package (pip install python-docx)",
            ) from err

        try:
            document = Document(io.BytesIO(raw))
        except Exception as err:  # noqa: BLE001 - docx raises many parse error types
            raise LoaderError("PARSE_ERROR", f"could not parse Word document: {err}") from err

        lines: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            heading = _heading_prefix(paragraph.style.name if paragraph.style else "")
            lines.append(f"{heading}{text}")

        for table in document.tables:
            lines.append("")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))

        text = "\n".join(lines)
        title = document.core_properties.title or os.path.splitext(os.path.basename(filename))[0]
        return [
            LoadedDocument(
                title=title.strip(),
                doc_type="POLICY",
                text=text,
                source_format="DOCX",
                ingestion_ref=filename,
            )
        ]


def _heading_prefix(style_name: str) -> str:
    lowered = style_name.lower()
    if lowered == "title":
        return "# "
    if lowered.startswith("heading"):
        try:
            level = int(style_name.split()[-1])
        except ValueError:
            level = 1
        return "#" * min(level, 6) + " "
    return ""